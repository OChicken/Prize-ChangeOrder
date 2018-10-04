#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Date    : 2018-09-27 21:27:16
# @Author  : Ma Seoyin (Ma.Seoyin@gmail.com)
# @Link    : https://github.com/OChicken
# @Version : V2

import os
import time
import docx
import xlrd
import win32com.client
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import comtypes.client
from shutil import copyfile


def Printed_Name(Name):
    output = ''
    for i in range(0, len(Name)):
        if i != (len(Name) - 1):
            output = output + Name[i] + '、'
        else:
            output = output + Name[i] + '  同学：'
    return output


def Exchange_NameOrder(Name):
    Name.append(Name[0])
    del Name[0]
    return Printed_Name(Name)


def delEmptyElement(List):
    while List[-1] == '':
        del List[-1]
    return List


def WriteDocxFiles(template, Prize_th, NameDisplay, k):
    template.paragraphs[10 * k].clear()
    run = template.paragraphs[10 * k].add_run(NameDisplay)
    run.font.size = Pt(22)
    run.font.name = u'华文新魏'
    run.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'华文新魏')
    template.paragraphs[10 * k + 4].clear()
    run = template.paragraphs[10 * k + 4].add_run(Prize_th)
    run.font.size = Pt(26)
    run.font.name = u'华文新魏'
    run.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'华文新魏')


# 欢迎辞
Dir = os.getcwd() + '/'
word = comtypes.client.CreateObject('Word.Application')
doc_app=win32com.client.Dispatch('Word.Application')
print('欢迎使用 物理学术竞赛智育分奖状生成 小程序 :)\n'
      '<<<<<<<<<<<<<<<<<<<<< 啦啦啦我是分割线 >>>>>>>>>>>>>>>>>>>>>\n'
      'author:\n马守然 (2014级应用物理学)\n学术科创部\n物理与光电学院团委学生会\nEmail: 1941688873@qq.com / Ma.Seoyin@gmail.com\n'
      'Link: https://github.com/OChicken\n'
      '<<<<<<<<<<<<<<<<<<<<< 啦啦啦我是分割线 >>>>>>>>>>>>>>>>>>>>>\n'
      '奖状生成中, 请稍候......')
Start = time.clock()
PrizeList = xlrd.open_workbook(Dir + 'Prize.xlsx')
PrizeChinese = ['一等奖', '二等奖', '三等奖']
for i in range(3):
    # 新建第几等奖的文件夹
    PrizeDir = Dir + PrizeChinese[i] + '/'
    if os.path.exists(PrizeDir.rstrip('/')) == False:
        os.makedirs(PrizeDir)
    PrizeTemplate = PrizeDir + 'template.docx'
    copyfile('template.docx', PrizeTemplate)
    # 第几等奖
    Prize_th = PrizeChinese[i]
    # 第几等奖的所有队伍
    Prize = PrizeList.sheet_by_index(i)
    # 第几等奖的第j个队伍
    for j in range(Prize.ncols):
        Team = str(int(Prize.col_values(j)[0]))  # 的队号
        Name = Prize.col_values(j)[1:]  # 和名字序列
        delEmptyElement(Name)
        size = len(Name)
        # 生成名字序列那么长的奖状页数 (譬如说一个队有5人, 生成的奖状页数就是5页)
        for k in range(size - 1):
            if k == 0:
                doc=doc_app.Documents.Open(Dir + 'template.docx')
                doc.Content.Copy()
                sel = doc_app.Selection
                sel.Range.Paste()
                doc.SaveAs(PrizeTemplate)
                doc.Close()
            else:
                doc=doc_app.Documents.Open(PrizeTemplate)
                sel = doc_app.Selection
                sel.Range.Paste()
                doc.SaveAs(PrizeTemplate)
                doc.Close()
        template = docx.Document(PrizeTemplate)
        Leader = Name[0]  # 队长名字
        FileName = PrizeDir + Team + Leader
        for k in range(size):
            if k == 0:
                NameDisplay = Printed_Name(Name)
                WriteDocxFiles(template, Prize_th, NameDisplay, k)
            else:
                NameDisplay = Exchange_NameOrder(Name)
                WriteDocxFiles(template, Prize_th, NameDisplay, k)
        template.save(FileName + '.docx')
        doc = word.Documents.Open(FileName + '.docx')
        doc.SaveAs(FileName + '.pdf', FileFormat=17)
        doc.Close()
        os.remove(PrizeTemplate)
End = time.clock()
print('所有奖状已生成, 用时' + str(End-Start) + '秒')
input('请按回车关闭本宝宝 :)')
